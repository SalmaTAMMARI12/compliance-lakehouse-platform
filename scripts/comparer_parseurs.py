"""
Benchmark multi-parseurs — compare la qualité d'extraction de tableaux
sur un rapport d'audit DGSSI/PASSI (PDF ou DOCX).

Parseurs testés (selon disponibilité) :
  - Docling         (actuel, baseline)
  - pdfplumber      (PDF uniquement)
  - pymupdf/fitz    (PDF uniquement)
  - python-docx     (DOCX uniquement)

Usage :
    venv/Scripts/python scripts/comparer_parseurs.py <chemin_fichier>

Exemple :
    venv/Scripts/python scripts/comparer_parseurs.py "data/private/rapport2.pdf"
    venv/Scripts/python scripts/comparer_parseurs.py "data/private/rapport2.docx"

Sortie : console + fichier JSON de rapport dans le répertoire courant.
"""
from __future__ import annotations

import json
import sys
import time
from pathlib import Path

# ---------------------------------------------------------------------------
# Helpers de métriques
# ---------------------------------------------------------------------------

RICH_CELL_MARKER = "<!-- rich cell -->"


def _metriques_tableaux(tableaux: list[list[list[str]]]) -> dict:
    """Calcule des métriques de qualité sur une liste de tableaux extraits."""
    if not tableaux:
        return {
            "nb_tableaux": 0,
            "nb_cellules_total": 0,
            "nb_rich_cells": 0,
            "nb_cellules_vides": 0,
            "nb_cellules_utiles": 0,
            "pct_rich_cells": 0.0,
            "pct_cellules_utiles": 0.0,
            "exemples_cellules_utiles": [],
        }

    nb_total = 0
    nb_rich = 0
    nb_vides = 0
    exemples: list[str] = []

    for tableau in tableaux:
        for ligne in tableau:
            for cellule in ligne:
                nb_total += 1
                val = str(cellule).strip()
                if val == RICH_CELL_MARKER:
                    nb_rich += 1
                elif not val:
                    nb_vides += 1
                elif len(exemples) < 8:
                    exemples.append(val[:80])

    nb_utiles = nb_total - nb_rich - nb_vides
    return {
        "nb_tableaux": len(tableaux),
        "nb_cellules_total": nb_total,
        "nb_rich_cells": nb_rich,
        "nb_cellules_vides": nb_vides,
        "nb_cellules_utiles": nb_utiles,
        "pct_rich_cells": round(nb_rich / nb_total * 100, 1) if nb_total else 0.0,
        "pct_cellules_utiles": round(nb_utiles / nb_total * 100, 1) if nb_total else 0.0,
        "exemples_cellules_utiles": exemples,
    }


def _scorer(metriques: dict) -> float:
    """Score simple 0–100 : plus c'est élevé, meilleure est la qualité."""
    return metriques["pct_cellules_utiles"]


# ---------------------------------------------------------------------------
# Parseur 1 : Docling (baseline actuelle)
# ---------------------------------------------------------------------------

def _essayer_docling(chemin: Path) -> dict:
    debut = time.perf_counter()
    try:
        from docling.datamodel.base_models import ConversionStatus, InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.document_converter import DocumentConverter, PdfFormatOption

        options = PdfPipelineOptions()
        options.do_ocr = False
        conv = DocumentConverter(
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=options)}
        )
        resultat = conv.convert(str(chemin))
        duree = time.perf_counter() - debut

        if resultat.status != ConversionStatus.SUCCESS:
            return {"statut": "ECHEC", "erreur": str(resultat.errors), "duree_s": round(duree, 1)}

        document = resultat.document
        tableaux = [
            table.export_to_dataframe(doc=document).values.tolist()
            for table in document.tables
        ]
        m = _metriques_tableaux(tableaux)
        return {"statut": "OK", "duree_s": round(duree, 1), "score": _scorer(m), **m}

    except Exception as e:
        return {"statut": "ERREUR", "erreur": str(e), "duree_s": round(time.perf_counter() - debut, 1)}


# ---------------------------------------------------------------------------
# Parseur 2 : pdfplumber (PDF uniquement)
# ---------------------------------------------------------------------------

def _essayer_pdfplumber(chemin: Path) -> dict:
    debut = time.perf_counter()
    try:
        import pdfplumber  # type: ignore

        tableaux: list[list[list[str]]] = []
        with pdfplumber.open(str(chemin)) as pdf:
            for page in pdf.pages:
                for table in page.extract_tables():
                    if table:
                        # pdfplumber retourne None pour les cellules vides
                        tableau_propre = [
                            [str(cell).strip() if cell is not None else "" for cell in ligne]
                            for ligne in table
                        ]
                        tableaux.append(tableau_propre)

        duree = time.perf_counter() - debut
        m = _metriques_tableaux(tableaux)
        return {"statut": "OK", "duree_s": round(duree, 1), "score": _scorer(m), **m}

    except ImportError:
        return {"statut": "NON_INSTALLE", "erreur": "pip install pdfplumber",
                "duree_s": round(time.perf_counter() - debut, 1)}
    except Exception as e:
        return {"statut": "ERREUR", "erreur": str(e), "duree_s": round(time.perf_counter() - debut, 1)}


# ---------------------------------------------------------------------------
# Parseur 3 : pymupdf / fitz (PDF uniquement)
# ---------------------------------------------------------------------------

def _essayer_pymupdf(chemin: Path) -> dict:
    debut = time.perf_counter()
    try:
        import fitz  # type: ignore  (pymupdf)

        tableaux: list[list[list[str]]] = []
        doc = fitz.open(str(chemin))
        for page in doc:
            for table in page.find_tables():
                lignes = []
                for ligne in table.extract():
                    lignes.append([str(c).strip() if c is not None else "" for c in ligne])
                if lignes:
                    tableaux.append(lignes)
        doc.close()

        duree = time.perf_counter() - debut
        m = _metriques_tableaux(tableaux)
        return {"statut": "OK", "duree_s": round(duree, 1), "score": _scorer(m), **m}

    except ImportError:
        return {"statut": "NON_INSTALLE", "erreur": "pip install pymupdf",
                "duree_s": round(time.perf_counter() - debut, 1)}
    except Exception as e:
        return {"statut": "ERREUR", "erreur": str(e), "duree_s": round(time.perf_counter() - debut, 1)}


# ---------------------------------------------------------------------------
# Parseur 4 : python-docx (DOCX uniquement)
# ---------------------------------------------------------------------------

def _essayer_python_docx(chemin: Path) -> dict:
    debut = time.perf_counter()
    try:
        from docx import Document  # type: ignore

        doc = Document(str(chemin))
        tableaux: list[list[list[str]]] = []
        for table in doc.tables:
            lignes = []
            for row in table.rows:
                lignes.append([cell.text.strip() for cell in row.cells])
            if lignes:
                tableaux.append(lignes)

        duree = time.perf_counter() - debut
        m = _metriques_tableaux(tableaux)
        return {"statut": "OK", "duree_s": round(duree, 1), "score": _scorer(m), **m}

    except ImportError:
        return {"statut": "NON_INSTALLE", "erreur": "pip install python-docx",
                "duree_s": round(time.perf_counter() - debut, 1)}
    except Exception as e:
        return {"statut": "ERREUR", "erreur": str(e), "duree_s": round(time.perf_counter() - debut, 1)}


# ---------------------------------------------------------------------------
# Affichage
# ---------------------------------------------------------------------------

def _afficher_resultat(nom: str, res: dict) -> None:
    statut_label = {"OK": "[OK]", "ECHEC": "[ECHEC]", "ERREUR": "[ERREUR]", "NON_INSTALLE": "[NON INSTALLE]",}.get(
        res.get("statut", ""), "[?]"
    )
    print("\n" + "-" * 60)
    print(f"  {statut_label}  {nom}")
    print("-" * 60)

    if res.get("statut") != "OK":
        print(f"  Statut  : {res.get('statut')}")
        print(f"  Erreur  : {res.get('erreur', 'inconnue')}")
        print(f"  Duree   : {res.get('duree_s', '?')} s")
        return

    score = res.get("score", 0)
    nb_blocs = int(score / 5)
    barre = "|" * nb_blocs + "." * (20 - nb_blocs)
    print(f"  Score            : {score:.1f}% [{barre}]")
    print(f"  Duree            : {res['duree_s']} s")
    print(f"  Tableaux trouves : {res['nb_tableaux']}")
    print(f"  Cellules totales : {res['nb_cellules_total']}")
    print(f"  Rich cells       : {res['nb_rich_cells']}  ({res['pct_rich_cells']}%)")
    print(f"  Cellules vides   : {res['nb_cellules_vides']}")
    print(f"  Cellules utiles  : {res['nb_cellules_utiles']}  ({res['pct_cellules_utiles']}%)")
    if res.get("exemples_cellules_utiles"):
        print("  Exemples extraits :")
        for ex in res["exemples_cellules_utiles"][:5]:
            print(f"    >> {ex!r}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def comparer(chemin_fichier: Path) -> None:
    # Vérification immédiate — évite 87 secondes de timeout Docling inutiles
    if not chemin_fichier.exists():
        print(f"\n ERREUR : Fichier introuvable -> {chemin_fichier}")
        print(" Verifiez le chemin et relancez.")
        print(f"\n Exemple :")
        print(f"   venv\\Scripts\\python scripts\\comparer_parseurs.py \"C:\\Users\\hp\\Desktop\\monrapport.pdf\"")
        sys.exit(1)

    ext = chemin_fichier.suffix.lower()
    print(f"\n{'='*60}")
    print(f"  COMPARAISON DE PARSEURS")
    print(f"  Fichier : {chemin_fichier.name}")
    print(f"  Format  : {ext.upper()}")
    print(f"{'='*60}")

    resultats: dict[str, dict] = {}

    if ext == ".pdf":
        print("\n[1/3] Docling (baseline actuelle)...")
        resultats["Docling"] = _essayer_docling(chemin_fichier)

        print("[2/3] pdfplumber...")
        resultats["pdfplumber"] = _essayer_pdfplumber(chemin_fichier)

        print("[3/3] pymupdf (fitz)...")
        resultats["pymupdf"] = _essayer_pymupdf(chemin_fichier)

    elif ext in (".docx", ".doc"):
        print("\n[1/2] python-docx (natif Word)...")
        resultats["python-docx"] = _essayer_python_docx(chemin_fichier)

        print("[2/2] Docling (baseline actuelle)...")
        resultats["Docling"] = _essayer_docling(chemin_fichier)

    else:
        print(f"Format '{ext}' non supporté. Fichiers acceptés : .pdf, .docx")
        sys.exit(1)

    # Affichage des résultats individuels
    for nom, res in resultats.items():
        _afficher_resultat(nom, res)

    # Classement final
    print(f"\n{'='*60}")
    print("  CLASSEMENT FINAL (par score de qualité)")
    print(f"{'='*60}")
    ok_resultats = [(n, r) for n, r in resultats.items() if r.get("statut") == "OK"]
    ok_resultats.sort(key=lambda x: x[1].get("score", 0), reverse=True)

    for rang, (nom, res) in enumerate(ok_resultats, 1):
        medaille = ["#1", "#2", "#3"][rang - 1] if rang <= 3 else f"#{rang}"
        print(f"  [{medaille}]  {nom:15s}  score={res['score']:.1f}%  duree={res['duree_s']}s  "
              f"rich_cells={res['pct_rich_cells']}%")

    if ok_resultats:
        gagnant = ok_resultats[0][0]
        print(f"\n  >> Recommandation : utiliser [{gagnant}] pour ce rapport")
    else:
        print("\n  /!\ Aucun parseur n'a reussi. Verifiez le chemin du fichier.")

    # Sauvegarde JSON — dans le dossier scripts/ (toujours accessible en ecriture)
    rapport_json = {
        "fichier": str(chemin_fichier),
        "format": ext,
        "resultats": resultats,
        "classement": [n for n, _ in ok_resultats],
    }
    dossier_sortie = Path(__file__).parent
    chemin_json = dossier_sortie / f"benchmark_{chemin_fichier.stem}.json"
    try:
        chemin_json.write_text(json.dumps(rapport_json, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"\n  Rapport JSON sauvegarde : {chemin_json}")
    except Exception as e:
        print(f"\n  (Sauvegarde JSON ignoree : {e})")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python scripts/comparer_parseurs.py <chemin_fichier.pdf|.docx>")
        sys.exit(1)
    comparer(Path(sys.argv[1]))
